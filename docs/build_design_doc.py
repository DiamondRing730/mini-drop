"""Build the submission-ready Mini-Drop design document.

Run with the bundled workspace Python so python-docx/Pillow versions are stable.
"""
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Mini-Drop设计文档.docx"
ASSETS = ROOT / "design-assets"
ASSETS.mkdir(exist_ok=True)

NAVY = "16324F"
BLUE = "2563A6"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "16784A"
RED = "B42318"
GOLD = "A15C00"
WHITE = "FFFFFF"
BLACK = "20242A"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def pc(hex_value: str) -> str:
    """Convert an OOXML six-digit color to Pillow's #RRGGBB form."""
    return hex_value if hex_value.startswith("#") else f"#{hex_value}"


def set_font(run, size=11, bold=False, color=BLACK, italic=False, ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_font(run, size=9.2, bold=True, color=WHITE)
    set_repeat_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_idx % 2 == 1:
                shade_cell(cells[i], "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items, compact=True):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3 if compact else 6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        set_font(r, size=10.2)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        set_font(r, size=10.2)


def add_callout(doc, label, text, tone="blue"):
    fills = {"blue": LIGHT_BLUE, "gray": LIGHT_GRAY, "green": "EAF6EF", "gold": "FFF5E6"}
    colors = {"blue": BLUE, "gray": MID_GRAY, "green": GREEN, "gold": GOLD}
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fills[tone])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}  ")
    set_font(lead, size=10.2, bold=True, color=colors[tone])
    body = p.add_run(text)
    set_font(body, size=10.2, color=BLACK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "111827")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=8.8, color="E5E7EB", ascii_font="Consolas")
    set_table_geometry(table, [9360])


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=8.5, color=MID_GRAY, italic=True)


def add_picture(doc, path, width, alt_text):
    doc.add_picture(str(path), width=width)
    shape = doc.inline_shapes[-1]
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)


def page_title(doc, number, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}  {title}")
    set_font(r, size=18, bold=True, color=NAVY)
    if subtitle:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(10)
        s = q.add_run(subtitle)
        set_font(s, size=10, color=MID_GRAY)


def page_break(doc):
    doc.add_page_break()


def pil_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def center_text(draw, box, text, font, fill):
    x1, y1, x2, y2 = box
    bb = draw.multiline_textbbox((0, 0), text, font=font, spacing=4, align="center")
    width, height = bb[2] - bb[0], bb[3] - bb[1]
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), text,
                        font=font, fill=pc(fill), spacing=4, align="center")


def arrow(draw, start, end, fill=BLUE, width=7):
    draw.line([start, end], fill=pc(fill), width=width)
    x, y = end
    sx, sy = start
    if abs(x - sx) > abs(y - sy):
        sign = 1 if x > sx else -1
        pts = [(x, y), (x - sign * 18, y - 12), (x - sign * 18, y + 12)]
    else:
        sign = 1 if y > sy else -1
        pts = [(x, y), (x - 12, y - sign * 18), (x + 12, y - sign * 18)]
    draw.polygon(pts, fill=pc(fill))


def architecture_image(path):
    img = Image.new("RGB", (1600, 820), pc(WHITE))
    d = ImageDraw.Draw(img)
    title = pil_font(38, True)
    label = pil_font(28, True)
    small = pil_font(22)
    d.text((60, 35), "Mini-Drop Runtime Architecture", font=title, fill=pc(NAVY))
    boxes = {
        "web": (70, 180, 350, 350),
        "server": (520, 150, 890, 380),
        "agent": (1050, 100, 1510, 310),
        "analyzer": (1050, 490, 1510, 700),
        "pg": (520, 520, 890, 700),
    }
    colors = {"web": "E8F1FA", "server": "DCEBFA", "agent": "EAF6EF", "analyzer": "FFF5E6", "pg": "F2F4F7"}
    content = {
        "web": ("WEB", "React + TypeScript\nTask UI / charts"),
        "server": ("SERVER", "FastAPI + SQLAlchemy\nOrchestration / state machine"),
        "agent": ("AGENT", "Host PID namespace\nperf / py-spy / bpftrace"),
        "analyzer": ("ANALYZER", "Folded stacks -> SVG\nTopN / eBPF distribution"),
        "pg": ("POSTGRESQL", "Tasks / transitions\nAgents / chunks / audit"),
    }
    for key, box in boxes.items():
        d.rounded_rectangle(box, radius=22, fill=pc(colors[key]), outline=pc(BLUE if key != "pg" else MID_GRAY), width=4)
        head, body = content[key]
        center_text(d, (box[0], box[1] + 16, box[2], box[1] + 75), head, label, NAVY)
        center_text(d, (box[0] + 15, box[1] + 80, box[2] - 15, box[3] - 15), body, small, BLACK)
    arrow(d, (350, 265), (520, 265))
    d.text((375, 220), "REST / polling", font=small, fill=pc(MID_GRAY))
    arrow(d, (890, 225), (1050, 205))
    d.text((895, 155), "heartbeat + task", font=small, fill=pc(MID_GRAY))
    arrow(d, (700, 380), (700, 520))
    d.text((720, 430), "SQL", font=small, fill=pc(MID_GRAY))
    arrow(d, (1280, 310), (1280, 490))
    d.text((1300, 385), "artifacts", font=small, fill=pc(MID_GRAY))
    arrow(d, (1050, 595), (890, 595))
    d.text((900, 550), "analysis result", font=small, fill=pc(MID_GRAY))
    img.save(path)


def flow_image(path):
    img = Image.new("RGB", (1600, 500), pc(WHITE))
    d = ImageDraw.Draw(img)
    label = pil_font(24, True)
    small = pil_font(20)
    steps = [
        ("1", "Create task"), ("2", "Heartbeat claim"), ("3", "Real profiling"),
        ("4", "Store artifact"), ("5", "Analyze"), ("6", "Render in Web"),
    ]
    x = 45
    for idx, (num, text) in enumerate(steps):
        box = (x, 150, x + 210, 315)
        d.rounded_rectangle(box, radius=18, fill=pc(LIGHT_BLUE if idx < 3 else "EAF6EF"), outline=pc(BLUE), width=3)
        d.ellipse((x + 75, 90, x + 135, 150), fill=pc(BLUE))
        center_text(d, (x + 75, 90, x + 135, 150), num, label, WHITE)
        center_text(d, box, text, label, NAVY)
        if idx < len(steps) - 1:
            arrow(d, (x + 210, 232), (x + 250, 232), width=5)
        x += 260
    d.text((55, 380), "All task state changes are persisted with a human-readable reason.", font=small, fill=pc(MID_GRAY))
    img.save(path)


def state_image(path):
    img = Image.new("RGB", (1600, 620), pc(WHITE))
    d = ImageDraw.Draw(img)
    font = pil_font(25, True)
    small = pil_font(20)
    boxes = {
        "PENDING": (60, 90, 325, 210),
        "RUNNING": (430, 90, 695, 210),
        "UPLOADING": (800, 90, 1100, 210),
        "DONE": (1210, 90, 1500, 210),
        "FAILED": (610, 390, 900, 510),
        "STOPPED": (1030, 390, 1320, 510),
    }
    fills = {"DONE": "EAF6EF", "FAILED": "FDECEC", "STOPPED": "F2ECFF"}
    for name, box in boxes.items():
        fill = fills.get(name, LIGHT_BLUE)
        outline = GREEN if name == "DONE" else RED if name == "FAILED" else BLUE
        d.rounded_rectangle(box, radius=20, fill=pc(fill), outline=pc(outline), width=4)
        center_text(d, box, name, font, NAVY if name not in ("FAILED",) else RED)
    arrow(d, (325, 150), (430, 150))
    arrow(d, (695, 150), (800, 150))
    arrow(d, (1100, 150), (1210, 150))
    for x in (190, 560, 950):
        arrow(d, (x, 210), (730, 390), fill=RED, width=4)
    arrow(d, (560, 210), (1160, 390), fill="7047A5", width=4)
    d.text((65, 545), "Terminal states: DONE / FAILED / STOPPED. Continuous sessions stop at the next slice boundary.", font=small, fill=pc(MID_GRAY))
    img.save(path)


def optimization_image(path):
    img = Image.new("RGB", (1400, 560), pc(WHITE))
    d = ImageDraw.Draw(img)
    title = pil_font(32, True)
    label = pil_font(23)
    value = pil_font(25, True)
    d.text((55, 30), "Verified hotspot change: fib self-time share", font=title, fill=pc(NAVY))
    axis_x, axis_y, axis_w = 310, 420, 960
    d.line((axis_x, axis_y, axis_x + axis_w, axis_y), fill=pc(MID_GRAY), width=3)
    for tick in range(0, 81, 20):
        x = axis_x + int(axis_w * tick / 80)
        d.line((x, axis_y, x, axis_y + 12), fill=pc(MID_GRAY), width=2)
        d.text((x - 15, axis_y + 20), str(tick), font=label, fill=pc(MID_GRAY))
    rows = [("Before: naive recursion", 74.22, RED, 150), ("After: lru_cache", 0.0, GREEN, 290)]
    for name, pct, color, y in rows:
        d.text((55, y + 18), name, font=label, fill=pc(BLACK))
        width = max(3, int(axis_w * pct / 80))
        d.rounded_rectangle((axis_x, y, axis_x + width, y + 70), radius=12, fill=pc(color))
        d.text((axis_x + width + 18, y + 18), f"{pct:.2f}%", font=value, fill=pc(color))
    d.text((950, 500), "Independent recomputation: 5 / 5 checks passed", font=label, fill=pc(GREEN))
    img.save(path)


ARCH = ASSETS / "architecture.png"
FLOW = ASSETS / "flow.png"
STATE = ASSETS / "state-machine.png"
OPT = ASSETS / "optimization-proof.png"
architecture_image(ARCH)
flow_image(FLOW)
state_image(STATE)
optimization_image(OPT)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.76)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Preset: standard_business_brief. Named layout override: 0.82 in side/top margins
# and 0.76 in bottom margin to keep the required report within nine pages.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = rgb(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

style_tokens = {
    "Heading 1": (16, BLUE, 16, 8),
    "Heading 2": (13, BLUE, 12, 6),
    "Heading 3": (12, NAVY, 8, 4),
}
for name, (size, color, before, after) in style_tokens.items():
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    style = doc.styles[list_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.2)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run("MINI-DROP  |  Linux 性能诊断系统设计文档")
set_font(hr, size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("DiamondRing730  |  ")
set_font(fr, size=8.2, color=MID_GRAY)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Page 1 — memo masthead / executive summary.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(34)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("MINI-DROP")
set_font(r, size=30, bold=True, color=NAVY)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(12)
r = p2.add_run("Linux 按需性能采集、可视化与可验证优化平台")
set_font(r, size=16, bold=True, color=BLUE)

add_table(doc, ["文档属性", "内容"], [
    ("版本", "1.0 · 最终交付版 · 2026-06-21"),
    ("代码基线", "main / 7b3ea88 · github.com/DiamondRing730/mini-drop"),
    ("验证环境", "WSL2 Ubuntu 22.04 · Linux 6.6 · Docker Desktop"),
], [1900, 7460], font_size=9.5)

add_heading(doc, "执行摘要", 1)
add_body(doc, "Mini-Drop 复刻真实性能诊断系统的核心闭环：用户在 Web 选择容器和进程，下发采样任务；Server 通过持久化状态机编排任务；Agent 调用 perf、py-spy 或 bpftrace 真实采集；Analyzer 生成火焰图、热点与延迟分布；Web 展示结果并支持离线/DeepSeek 归因和优化前后验证。")
add_callout(doc, "完成度", "四组件 + PostgreSQL 可由 Docker Compose 一键启动；46 条单测、3 条 E2E 通过；端到端、eBPF、Continuous Profiling 和优化闭环均已在 WSL2 实跑。", "green")
add_heading(doc, "设计目标", 2)
add_bullets(doc, [
    "真实性：采样、状态迁移、产物和分析结果全部来自运行时，不使用预生成演示数据。",
    "完整性：覆盖目标发现、调度、采集、分析、可视化、归因、复测和异常恢复。",
    "可验证：每次状态迁移带 reason；AI 结论和优化差分均由独立逻辑复算。",
    "可复现：评审仅需 docker compose up 和 make demo，即可在 10 分钟内完成演示。",
])
add_heading(doc, "文档导航", 2)
add_body(doc, "第 2–3 页给出架构与状态机；第 4–6 页描述组件、采集、分析和创新闭环；第 7–8 页记录关键决策、可靠性与性能自证；第 9 页总结 AI 协作及未来 7 天计划。")

page_break(doc)

# Page 2 — requirements and architecture.
page_title(doc, "01", "需求映射与总体架构", "以真实跑通、功能覆盖和可验证结果为第一优先级")
add_table(doc, ["目标能力", "实现", "验收信号"], [
    ("按需采集", "perf / py-spy / eBPF 三类采集器", "真实目标 PID，结果可下载"),
    ("状态机", "PENDING→RUNNING→UPLOADING→DONE/FAILED；持续任务含 STOPPED", "迁移落库并带 reason"),
    ("Agent 管理", "5 秒心跳、30 秒离线、原子领取、重启接管", "REGISTER/OFFLINE/RECOVER 审计"),
    ("持续画像", "py-spy 定时切片、时间轴、窗口合并、停止/续建", "任意窗口在线渲染火焰图"),
    ("智能分析", "离线规则或显式 DeepSeek；只读工具 + 独立校验", "结论、证据、建议、校验率"),
], [1900, 4700, 2760], font_size=8.8)
add_picture(doc, ARCH, Inches(6.55), "Mini-Drop runtime architecture with Web, Server, Agent, Analyzer and PostgreSQL")
caption(doc, "图 1  运行时架构：控制面与采集/分析面解耦，产物通过共享卷交接")
add_body(doc, "Server 是唯一编排中心；Agent 运行于 host PID namespace 并拥有采集权限；Analyzer 不执行采集，只消费已落盘产物。该边界让采集工具、分析算法和前端展示可以独立演进。")

page_break(doc)

# Page 3 — flow and state machine.
page_title(doc, "02", "端到端数据流与状态机", "一条任务从用户操作到可视化结果的完整生命线")
add_picture(doc, FLOW, Inches(6.55), "End-to-end task flow from creation through profiling and analysis to Web rendering")
caption(doc, "图 2  端到端数据流：Web 下发 → 心跳领取 → 真实采集 → 分析 → 展示")
add_picture(doc, STATE, Inches(6.55), "Task state machine with DONE, FAILED and STOPPED terminal states")
caption(doc, "图 3  任务状态机：终态不可原地重试，续建和重试均创建新任务")
add_body(doc, "Server 使用 FOR UPDATE SKIP LOCKED 原子领取 PENDING 任务，防止多个 Agent 重复执行。所有状态变化由统一 transition() 校验并写入 task_state_transitions；非法跳转返回 409。运行中任务禁止删除，持续任务通过心跳下发停止信号，在当前切片结束后进入 STOPPED。")
add_callout(doc, "故障恢复", "Agent 心跳线程与 worker 分离。单任务异常被最外层隔离；重启后 Agent 上报 active_task_ids，Server 自动重新下发数据库中仍为 RUNNING、但新进程未持有的任务。", "blue")

page_break(doc)

# Page 4 — component design.
page_title(doc, "03", "组件与接口设计", "四个进程各自保持单一职责，通过明确契约协作")
add_table(doc, ["组件", "核心职责", "关键实现"], [
    ("Web", "创建/筛选任务；状态与审计；火焰图、TopN、eBPF、时间轴、归因和差分", "React + TypeScript + ECharts；2–3 秒轮询"),
    ("Server", "任务编排、状态机、Agent 管理、切片索引、产物下载、归因与差分 API", "FastAPI + SQLAlchemy + PostgreSQL"),
    ("Agent", "目标发现、心跳领取、采集器调度、自监控、切片与结果上报", "privileged + pid:host；独立 worker"),
    ("Analyzer", "折叠栈解析、树构建、SVG 火焰图、TopN、eBPF 直方图", "纯 Python 分析流水线"),
    ("PostgreSQL", "任务、迁移、Agent、切片和生命周期审计", "5 张核心表；JSON 保存产物映射"),
], [1250, 4060, 4050], font_size=8.7)
add_heading(doc, "核心接口契约", 2)
add_table(doc, ["方向", "接口", "语义"], [
    ("Web→Server", "POST /api/v1/tasks", "创建一次性或持续任务"),
    ("Agent→Server", "POST /agent/heartbeat", "上报资源、容器/PID、active tasks；领取任务/停止信号"),
    ("Agent→Server", "POST /agent/tasks/{tid}/chunk", "登记持续采样切片"),
    ("Analyzer→Server", "GET /internal/analysis/next", "领取待分析产物"),
    ("Web→Server", "GET /tasks/{tid}/window", "合并重叠切片并实时返回 SVG"),
], [1700, 3500, 4160], font_size=8.7)
add_heading(doc, "数据模型", 2)
add_body(doc, "Task 保存目标、采样参数、状态、原因与产物映射；TaskStateTransition 形成不可丢失的状态审计；Agent 保存心跳、自监控和发现结果；ProfileChunk 保存切片起止时间、样本数和 folded 文件；AgentEvent 保存注册、离线与恢复事件。")

page_break(doc)

# Page 5 — collectors and continuous.
page_title(doc, "04", "采集系统与 Continuous Profiling", "按问题类型选择工具，而不是用一种采集器解释所有性能问题")
add_table(doc, ["采集器", "观察层级", "产物与用途"], [
    ("py-spy", "CPython 解释器帧", "folded stacks；Python 函数火焰图和 TopN"),
    ("perf", "原生 CPU 调用栈", "perf.data / perf script；系统与本地代码热点"),
    ("bpftrace", "内核 tracepoint", "read/write 延迟桶和进程计数；I/O 抖动定位"),
], [1600, 2800, 4960], font_size=9.2)
add_heading(doc, "目标发现", 2)
add_body(doc, "Agent 通过 Docker Unix socket 只发起 GET 请求，读取运行容器及 docker top 返回的宿主 PID。Web 以 Agent→容器→进程三级选择器填充目标 PID，同时保留手动输入。采集器仍以 /proc/{pid} 做最终存在性校验。")
add_heading(doc, "持续采样机制", 2)
add_numbered(doc, [
    "会话按 slice_sec（1–60 秒）循环调用 py-spy，每个切片输出独立 folded 文件。",
    "Agent 上报 start_ts、end_ts、samples 和文件名；Server 只索引元数据，不在热路径合并。",
    "用户选择时间窗口时，Server 合并所有与窗口相交的切片并即时渲染 SVG。",
    "主动停止通过心跳传递 Event，在当前切片结束后收口；续建创建新任务，旧时间轴不变。",
])
add_callout(doc, "边界", "当前 Continuous Profiling 是可停止的有限时长 py-spy 会话；窗口精度以切片为最小粒度。无限常驻、分层保留和 perf 持续模式留作后续。", "gold")
add_heading(doc, "采集安全", 2)
add_body(doc, "采集命令使用参数数组而非 shell 拼接，独立进程组用于超时清理；Agent 自监控 CPU、RSS 与 I/O，避免诊断工具本身成为不可见负载。")

page_break(doc)

# Page 6 — analysis, AI, verified optimization.
page_title(doc, "05", "分析、智能归因与可验证优化闭环", "把“哪里慢”推进到“改完是否真的更好”")
add_heading(doc, "分析流水线", 2)
add_body(doc, "Analyzer 将 py-spy/perf 折叠栈解析为调用树，计算函数自耗时 TopN，并用纯 Python 生成可内嵌 SVG；eBPF 原始输出被转换为延迟直方图与按进程分布。原始文件、tree.json、TopN 和 SVG 均作为任务产物保留。")
add_heading(doc, "受约束归因", 2)
add_body(doc, "页面显式选择离线或 DeepSeek。两种引擎共享只读 profile 工具；模型不能执行任意命令。独立 verifier 从 tree.json 复核函数存在性和自耗时比例，防止看似合理但与采样不一致的结论。离线评测 4 个确定性用例，Top-1、数值校验、建议覆盖和工具轨迹均为 100%。")
add_picture(doc, OPT, Inches(6.45), "Verified optimization chart showing fib self-time share decreasing from 74.22 percent to zero")
caption(doc, "图 4  可验证优化结果：相同 CPU 程序中 fib 热点占比由 74.22% 降至 0%")
add_body(doc, "优化闭环为：优化前采样 → 修改代码/配置 → 优化后采样 → 函数归一化 → 占比差分 → 红绿差分火焰图 → 独立复算。系统使用占比而非绝对样本数比较，并按最小样本量给出置信度。")
add_callout(doc, "不夸大结论", "热点占比下降证明 profile 分布变化，但不等价于整体延迟或吞吐必然提升；报告固定展示这一限制，并建议结合业务指标。", "gray")

page_break(doc)

# Page 7 — decisions, reliability, tradeoffs.
page_title(doc, "06", "关键决策、可靠性与取舍", "优先保证真实闭环，同时明确生产化缺口")
add_table(doc, ["决策", "为什么这样做", "代价 / 后续方向"], [
    ("FastAPI + Python Agent", "统一语言降低短周期实现成本；采集工具本身仍是真实 Linux 工具", "高并发控制面可迁移 Go；契约保持不变"),
    ("共享卷保存产物", "本机 Compose 最简单可靠，避免上传服务遮蔽核心链路", "多机改对象存储 + 校验和 + 生命周期"),
    ("轮询与心跳控制", "演示环境可观察、易调试，停止信号复用现有通道", "规模化改消息队列/长连接"),
    ("纯 Python 火焰图", "去除外部 Perl 脚本依赖，分析测试可重复", "超大 profile 需流式解析和缓存"),
    ("显式 AI 引擎选择", "默认离线且不上传；外部调用可控、可审计", "真实数据集与多轮稳定性仍需扩充"),
], [2100, 4000, 3260], font_size=8.4)
add_heading(doc, "可靠性机制", 2)
add_bullets(doc, [
    "状态机拒绝非法跳转；所有边带 reason；重试/续建创建新任务，保留历史证据链。",
    "FOR UPDATE SKIP LOCKED 保证原子领取；运行中删除返回 409，消除产物上报 404 竞态。",
    "心跳与 worker 分线程；最外层异常隔离；active_task_ids 支持 Agent 重启自动接管。",
    "30 秒无心跳标记 OFFLINE 并失败相关任务；恢复后写 RECOVER 审计。",
    "SVG 内联响应不设置下载头；产物路径 resolve 后必须位于任务根目录，防止目录穿越。",
])
add_callout(doc, "安全提示", "Docker socket 即使以 ro 文件方式挂载仍是高权限接口。当前发现模块仅实现 GET；生产环境应改用受限代理，只暴露容器/进程查询。", "gold")

page_break(doc)

# Page 8 — proof and reproducibility.
page_title(doc, "07", "性能自证、测试与可复现性", "用真实结果和自动化测试证明系统不是静态演示")
add_table(doc, ["证据", "结果", "说明"], [
    ("单元测试", "46 passed，约 80% 覆盖", "状态机、API、分析、eBPF、Continuous、归因、比较、恢复"),
    ("端到端", "3 passed", "正常采样、非法 PID、Agent 离线/恢复"),
    ("py-spy", "火焰图与 TopN 命中 fib", "端到端从下发到页面展示"),
    ("eBPF", "130,333 次 read/write；dd 7,470 次", "现场制造 I/O 后分布变化"),
    ("Continuous", "停止前 2 切片，最终保存 3 切片", "删除保护 409；停止进入 STOPPED"),
    ("优化闭环", "fib 74.22% → 0%；5/5 复算", "朴素递归对比 lru_cache"),
    ("归因评测", "4/4 Top-1；数值校验 100%", "确定性离线基线，可重复"),
], [2100, 2860, 4400], font_size=8.5)
add_heading(doc, "干净环境复现", 2)
add_code(doc, [
    "git clone https://github.com/DiamondRing730/mini-drop.git",
    "cd mini-drop",
    "docker compose up -d --build",
    "make demo        # 单场景端到端演示",
    "make unit        # 46 passed",
    "make e2e         # 3 passed；不可只看退出码或 skipped",
])
add_heading(doc, "环境与权限", 2)
add_body(doc, "验证环境为 WSL2 Ubuntu 22.04 / Linux 6.6 / amd64。Agent 需要 privileged、pid:host、SYS_PTRACE、SYS_ADMIN/PERFMON，并挂载 debugfs/tracefs；WSL 下 perf 使用 cpu-clock 软件事件规避硬件 PMU 限制。Web 暴露 8080，Server 暴露 8000。")
add_callout(doc, "验收原则", "评审应从新 clone 启动，而不是复用开发机数据库；E2E 输出必须明确为 3 passed，Server 不可达时的 skipped 不算通过。", "blue")

page_break(doc)

# Page 9 — AI collaboration and seven-day roadmap.
page_title(doc, "08", "AI 协作、未来 7 天与结论", "AI 加速实现，人负责目标、边界、实跑和证据")
add_heading(doc, "AI 协作方式", 2)
add_table(doc, ["AI 参与", "人工责任"], [
    ("拆解需求、生成初版代码、补测试、分析日志、整理文档", "确定功能优先级、选择环境、执行真实命令、检查页面和结果"),
    ("提出归因和优化建议", "要求工具受限、证据可见、数值独立复算，不接受无证据结论"),
    ("快速迭代故障修复", "提供实际输出并复测；例如发现 worker 线程因删除竞态退出"),
], [4680, 4680], font_size=8.8)
add_body(doc, "完整的代表性指令、AI 辅助内容和人工验收记录见 docs/AI_USAGE.md。仓库不保存 API Key、凭据或私密完整对话。")
add_heading(doc, "如果再有 7 天", 2)
add_numbered(doc, [
    "第 1–2 天：把共享卷替换为对象存储，增加 SHA-256 校验、断点上传和自动过期策略。",
    "第 3 天：使用受限 Docker API 代理，并按 Agent/用户建立最小权限与审计。",
    "第 4 天：Continuous Profiling 变为真正常驻服务，增加分层采样、压缩和跨会话时间轴。",
    "第 5 天：加入消息队列与幂等任务租约，支持多 Agent、取消确认、超时回收和水平扩展。",
    "第 6 天：扩充真实故障数据集，进行 DeepSeek 多轮稳定性、成本与错误类型评测。",
    "第 7 天：增加吞吐/延迟业务指标、OpenTelemetry 集成和 CI 干净机一键验收。",
])
add_heading(doc, "结论", 2)
add_body(doc, "Mini-Drop 已实现从目标发现、状态化调度、真实采集、分析可视化、可验证归因到优化复测的完整闭环。项目的核心价值不只是功能数量，而是每项能力均有运行时证据、失败路径和可复现测试；同时对共享卷、轮询、Docker 权限和有限时长 Continuous 等取舍保持透明。")
add_callout(doc, "最终交付", "Git 仓库保留完整提交历史；docker compose up + make demo 可运行；设计文档、AI 使用记录、归因评测报告与 ≤15 分钟演示视频共同构成交付证据。", "green")

# Metadata and save.
props = doc.core_properties
props.title = "Mini-Drop Linux 性能诊断系统设计文档"
props.subject = "腾讯 Mini 项目设计文档"
props.author = "DiamondRing730"
props.keywords = "Mini-Drop, performance profiling, eBPF, py-spy, Continuous Profiling"
props.comments = "Generated from the verified repository baseline 7b3ea88."

doc.save(OUT)
print(OUT)
